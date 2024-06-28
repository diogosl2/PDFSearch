import os
import tkinter as tk
from tkinter import filedialog
from openpyxl import Workbook
from openpyxl.styles import Alignment
import fitz  # PyMuPDF
import docx
from tqdm import tqdm

def buscar_termo_em_arquivos(termo, diretorio):
    arquivos_com_termo = []

    # Caminho dos arquivos no diretório selecionado
    arquivos = [os.path.join(root, file) for root, _, files in os.walk(diretorio) for file in files if file.lower().endswith((".pdf", ".docx", ".doc", ".txt"))]

    # Inicia a barra de progresso
    for file_path in tqdm(arquivos, desc="Buscando termos em arquivos"):
        try:
            # Verifica se o termo está presente no arquivo PDF, DOCX, DOC ou TXT
            if verificar_termo_no_arquivo(termo, file_path):
                count, linhas = termo_count_e_linhas_no_arquivo(termo, file_path)
                arquivos_com_termo.append((file_path, count, linhas))
        except Exception as e:
            print(f"Erro ao processar {file_path}: {e}")

    return arquivos_com_termo

def verificar_termo_no_arquivo(termo, file_path):
    try:
        if file_path.lower().endswith(".pdf"):
            return verificar_termo_no_pdf(termo, file_path)
        elif file_path.lower().endswith((".docx", ".doc")):
            return verificar_termo_no_docx(termo, file_path)
        elif file_path.lower().endswith(".txt"):
            return verificar_termo_no_txt(termo, file_path)
    except Exception as e:
        print(f"Erro ao verificar termo no arquivo {file_path}: {e}")
    return False

def verificar_termo_no_pdf(termo, file_path):
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if termo.lower() in text.lower():
                return True
        return False
    except Exception as e:
        print(f"Erro ao verificar termo no arquivo PDF {file_path}: {e}")
        return False

def verificar_termo_no_docx(termo, file_path):
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if termo.lower() in para.text.lower():
                return True
        return False
    except Exception as e:
        print(f"Erro ao verificar termo no arquivo DOCX {file_path}: {e}")
        return False

def verificar_termo_no_txt(termo, file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()
            if termo.lower() in text.lower():
                return True
        return False
    except Exception as e:
        print(f"Erro ao verificar termo no arquivo TXT {file_path}: {e}")
        return False

def termo_count_e_linhas_no_arquivo(termo, file_path):
    if file_path.lower().endswith(".pdf"):
        return termo_count_e_linhas_no_pdf(termo, file_path)
    elif file_path.lower().endswith((".docx", ".doc")):
        return termo_count_e_linhas_no_docx(termo, file_path)
    elif file_path.lower().endswith(".txt"):
        return termo_count_e_linhas_no_txt(termo, file_path)
    return 0, []

def termo_count_e_linhas_no_pdf(termo, file_path):
    count = 0
    linhas = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            lines = text.split('\n')
            for line in lines:
                if termo.lower() in line.lower():
                    linhas.append(line.strip())
                    count += 1
    except Exception as e:
        print(f"Erro ao contar termo no arquivo PDF {file_path}: {e}")
    return count, linhas

def termo_count_e_linhas_no_docx(termo, file_path):
    count = 0
    linhas = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if termo.lower() in para.text.lower():
                linhas.append(para.text.strip())
                count += 1
    except Exception as e:
        print(f"Erro ao contar termo no arquivo DOCX {file_path}: {e}")
    return count, linhas

def termo_count_e_linhas_no_txt(termo, file_path):
    count = 0
    linhas = []
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            lines = file.readlines()
            for line in lines:
                if termo.lower() in line.lower():
                    linhas.append(line.strip())
                    count += 1
    except Exception as e:
        print(f"Erro ao contar termo no arquivo TXT {file_path}: {e}")
    return count, linhas

def salvar_resultados_excel(arquivos_com_termo, termo):
    if not arquivos_com_termo:
        print(f"Nenhum arquivo encontrado com o termo '{termo}'.")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "Resultados de Busca"

    # Definindo estilos para as células
    alinhamento = Alignment(wrapText=True, horizontal="center", vertical="center")
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 20
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 80

    # Cabeçalho
    ws['A1'] = "Arquivo"
    ws['B1'] = "Tipo"
    ws['C1'] = "Termo Buscado"
    ws['D1'] = "Quantidade de Ocorrências"
    ws['E1'] = "Linhas Correspondentes"

    for row in ws.iter_rows(min_row=1, max_row=1, max_col=5):
        for cell in row:
            cell.alignment = alinhamento

    for arquivo, quantidade, linhas in arquivos_com_termo:
        arquivo_base = os.path.basename(arquivo)
        tipo_arquivo = obter_tipo_arquivo(arquivo)
        linhas_texto = "\n".join(linhas)
        ws.append([arquivo_base, tipo_arquivo, termo, quantidade, linhas_texto])

    nome_arquivo = f"resultados_busca_{termo}.xlsx"
    caminho_arquivo = os.path.join(os.path.dirname(os.path.abspath(__file__)), nome_arquivo)
    wb.save(caminho_arquivo)
    print(f"Resultados salvos em: {caminho_arquivo}")

def obter_tipo_arquivo(file_path):
    extensao = os.path.splitext(file_path)[1].lower()
    if extensao == ".pdf":
        return "PDF"
    elif extensao == ".docx" or extensao == ".doc":
        return "DOCX/DOC"
    elif extensao == ".txt":
        return "TXT"
    else:
        return "Desconhecido"

def selecionar_diretorio():
    root = tk.Tk()
    root.withdraw()
    diretorio = filedialog.askdirectory(title="Selecione o diretório com os arquivos")
    return diretorio

if __name__ == "__main__":
    diretorio = selecionar_diretorio()
    if not diretorio:
        print("Nenhum diretório selecionado. Encerrando o programa.")
        exit()

    termo = input("Digite o termo a ser buscado nos arquivos PDF, DOCX, DOC e TXT: ")

    arquivos_encontrados = buscar_termo_em_arquivos(termo, diretorio)

    if arquivos_encontrados:
        salvar_resultados_excel(arquivos_encontrados, termo)
    else:
        print(f"\nNenhum arquivo encontrado com o termo '{termo}'.")
